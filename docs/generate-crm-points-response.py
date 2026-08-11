#!/usr/bin/env python3
"""
Generates docs/CRM-Points-Response.docx - the client-facing response to the manager
review recorded in "CRM Points.docx".

Sources (all read before writing; this script only renders the conclusions):
  docs/superpowers/specs/2026-08-11-crm-points-manager-feedback-design.md
  docs/role-matrix.md
  docs/security-audit.md
  docs/scalability-and-storage.md
  the two feature commits on feat/crm-points-manager-feedback
  the source tree itself

Usage:  python docs/generate-crm-points-response.py
Requires: python-docx 1.1.2
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

OUT = Path(__file__).resolve().parent / "CRM-Points-Response.docx"

# --------------------------------------------------------------------------------------
# helpers
# --------------------------------------------------------------------------------------

doc = Document()


def _base_styles():
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Heading 1", 16, RGBColor(0x15, 0x37, 0x2B)),
        ("Heading 2", 13, RGBColor(0x1E, 0x5A, 0x3C)),
        ("Heading 3", 11.5, RGBColor(0x33, 0x33, 0x33)),
    ):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        st.paragraph_format.space_after = Pt(4)


def h1(text):
    return doc.add_heading(text, level=1)


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def para(text="", bold=False, italic=False, size=None, space_after=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def rich(parts, style=None):
    """parts = list of (text, {'b':True,'i':True,'mono':True})"""
    p = doc.add_paragraph(style=style)
    for text, fmt in parts:
        r = p.add_run(text)
        r.bold = bool(fmt.get("b"))
        r.italic = bool(fmt.get("i"))
        if fmt.get("mono"):
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    return p


def numbered(text):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.space_after = Pt(3)
    return p


def labelled(label, text):
    """Bold run-in label followed by body text, e.g. 'What was asked. ...'"""
    p = doc.add_paragraph()
    r = p.add_run(label + "  ")
    r.bold = True
    p.add_run(text)
    return p


def quote(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    return p


def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    return p


def table(headers, rows, widths=None, font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(htext)
        r.bold = True
        r.font.size = Pt(font_size)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val))
            r.font.size = Pt(font_size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def point(num, title):
    h1(f"P{num} — {title}")


# --------------------------------------------------------------------------------------
_base_styles()

# ======================================================================================
# Title
# ======================================================================================
t = doc.add_heading("CRM Review Points — Response", level=0)
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
para("Automation Systems NG Pvt Ltd — CRM (crm.automationsystems.info)", bold=True)
para("Response to the manager review recorded in CRM Points.docx")
para("Date: 11 August 2026     Branch: feat/crm-points-manager-feedback")
para(
    "Prepared from the source code as it stands on that branch. Every behaviour described "
    "below was read out of the code, not assumed.",
    italic=True,
)

# ======================================================================================
# Executive summary
# ======================================================================================
h1("Executive summary")

para(
    "All eleven points from your review document are implemented. A twelfth item — a "
    "performance bug found while doing the scalability analysis you asked for — was fixed "
    "with the owner's approval and is reported here as P12."
)

para("The short version:")
bullet(
    "The ownership model changed underneath. Case ownership used to be worked out fresh on "
    "every read from the customer's handler list. It is now stored on the case itself. That "
    "single change is what makes P1, P9, P10 and P11 possible at all; the old design could "
    "not express \"adding a handler adds an owner, but removing a handler does not remove "
    "one\", because there was nothing to remove from."
)
bullet(
    "L5 and L6 can no longer be account handlers, anywhere — not by hand, not from a bulk "
    "import, not as a silent default. They remain fully eligible as case owners and ticket "
    "assignees, which is what you asked for."
)
bullet(
    "The bug you spotted from the screenshot was real, and the root cause was worse than the "
    "symptom: one function was answering both \"who owns this case\" and \"is this person the "
    "account handler\", so the creator of a case on a Direct-handled account was labelled the "
    "account handler and could not be removed. Those two questions are now separate."
)
bullet(
    "Three analyses were produced: a role-by-action permission matrix traced line by line to "
    "the code, a database security review, and a storage and scalability projection. Summaries "
    "are in P2, P3 and P4 below; the full documents are in the repository."
)
bullet(
    "Three further pre-existing bugs were found incidentally and fixed. None were in your "
    "document. They are listed in their own section because you should know they existed."
)

para(
    "Two things to be clear about up front. First, nothing has been deployed: four database "
    "migrations are written, self-verifying and tested, but have not been applied to any "
    "database, including production. The deployment section explains the order they must run "
    "in and why. Second, there is an honest list of remaining limitations at the end. The most "
    "immediate one is that the SEI name list ships empty and needs an L6 to fill it in before "
    "the field does anything useful."
)

h2("Verification status")
table(
    ["Gate", "Result"],
    [
        ["TypeScript typecheck", "Clean"],
        ["Unit tests (vitest)", "240 / 240 passing"],
        ["End-to-end tests (Playwright)", "21 / 21 passing"],
        ["Production build (next build)", "Succeeds"],
        ["Database migrations 0005-0008", "Written and self-verifying — NOT YET APPLIED anywhere"],
    ],
    widths=[2.3, 4.2],
)

# ======================================================================================
# P1
# ======================================================================================
point(1, "L5 and L6 must not be account owners, but may be case owners")

labelled(
    "What was asked.",
    "\"L5 & L6 should not be account owner but can be case owner -> validate this behaviour.\"",
)

labelled(
    "What we found.",
    "Half of it was already true, and the important half was not. When an L5 or L6 created a "
    "customer without naming a handler, the system already recorded the virtual Direct account "
    "as the handler rather than the L5/L6 creator (customers/service.ts, creatorHandlerEmail). "
    "But api_addHandler performed no role check whatsoever, so an L5 or L6 could be added as a "
    "handler manually, by anyone entitled to add handlers. Nothing in the system prevented it. "
    "Bulk import was worse — see the separate bugs section.",
)

labelled(
    "What changed.",
    "addHandler now rejects any target whose role is L5 or L6, with a message that explains the "
    "alternative rather than just refusing: the user \"can still be added as a case owner or a "
    "ticket assignee\". The bulk importer applies the same rule to handler emails coming out of "
    "the spreadsheet. Case ownership and ticket assignment were deliberately left alone: "
    "addCaseOwner and assignTicket exclude only the virtual Direct account, never a real role. "
    "Migration 0006 deletes any L5/L6 handler rows that already exist in the database.",
)

labelled(
    "Current behaviour.",
    "An L5 or L6 can never be an account handler of a customer — the path is closed at creation, "
    "at manual add, at import, and retroactively by migration. They can be added as a case "
    "owner, and they can hold a ticket. Because L5 and L6 are level 4 or above, they still see "
    "every customer in the system regardless; their access is role-derived, not handler-derived.",
)

# ======================================================================================
# P2
# ======================================================================================
point(2, "Role-based permission matrix")

labelled(
    "What was asked.",
    "\"Make a role-based matrix of who can do what for similar scenarios.\"",
)

labelled(
    "What we found.",
    "The rules existed but had never been written down in one place, and they are not a single "
    "ladder. Three independent mechanisms decide whether an action is allowed: the user's role "
    "level (L1-L6), their access level to the specific customer (FULL / NAME / NONE, which "
    "depends on handler membership and location tags, not only on role), and case visibility "
    "(owner, assignee, or FULL access to the customer). Several actions are gated by only the "
    "second or third of those, with no role check at all. That is worth knowing before reading "
    "the table.",
)

labelled(
    "What changed.",
    "Nothing in the code — this is a documentation deliverable. It was written after the "
    "ownership rewrite so that it describes final behaviour. Every row in the full document is "
    "traced to a specific file and line number; where a rule could not be located in code it is "
    "marked UNVERIFIED rather than asserted. The full version is docs/role-matrix.md.",
)

h2("Primary action matrix")
para(
    "Yes = always allowed at that level. No = never. Cond. = allowed only if the stated "
    "condition also holds. \"FULL\" refers to the customer-access table that follows.",
    size=9.5,
)

table(
    ["Action", "L1", "L2", "L3", "L4", "L5", "L6", "Condition"],
    [
        ["Create customer", "No", "Yes", "Yes", "Yes", "Yes", "Yes",
         "L5/L6 creators are recorded as Direct, not themselves"],
        ["Edit customer details", "No", "Cond.", "Cond.", "Cond.", "Cond.", "Cond.",
         "FULL access to that customer"],
        ["Edit priority", "No", "Cond.", "Cond.", "Cond.", "Cond.", "Cond.", "FULL access and L2+"],
        ["Edit location / type / archive", "No", "No", "Cond.", "Cond.", "Cond.", "Cond.",
         "FULL access and L3+. Location can never be emptied"],
        ["Delete customer (to recycle bin)", "No", "No", "Cond.", "Cond.", "Cond.", "Cond.",
         "L3+, FULL access per row, and the customer must have no cases and no quotations"],
        ["Add / remove account handler", "No", "Cond.", "Cond.", "Cond.", "No", "No",
         "Caller is already a handler OR is L3+. Target must be active and must not be L5/L6"],
        ["Create case", "No", "Yes", "Yes", "Yes", "Cond.", "Cond.",
         "L2+ and FULL access. At L5+ an assignee must be named explicitly"],
        ["Change case stage", "Cond.", "Cond.", "Cond.", "Cond.", "Cond.", "Cond.",
         "Case visibility only — no role gate. Blocked once Won/Lost"],
        ["Set case outcome", "Cond.", "Cond.", "Cond.", "Cond.", "Cond.", "Cond.",
         "Case visibility only — no role gate"],
        ["Add / remove case owner", "Cond.", "Cond.", "Cond.", "Cond.", "Cond.", "Cond.",
         "Visible to the case AND (already an owner OR L4+). Cannot leave a case with zero owners"],
        ["Assign / reassign ticket", "Cond.", "Cond.", "Cond.", "Cond.", "Cond.", "Cond.",
         "Case visibility only. Blocked while the case is closed, including Hold"],
        ["Create or upload a quotation", "No", "Cond.", "Cond.", "Cond.", "Cond.", "Cond.",
         "L2+ and FULL access to the customer"],
        ["Set quote status, generate PDF, save to Drive", "Cond.", "Cond.", "Cond.", "Cond.",
         "Cond.", "Cond.", "FULL access only — no role floor on these"],
        ["View own dashboard", "Yes", "Yes", "Yes", "Yes", "Yes", "Yes", "Always"],
        ["View another user's dashboard", "No", "No", "Cond.", "Yes", "Yes", "Yes",
         "L3 may view an L2 who shares a location. L4+ may view anyone"],
        ["View the Direct dashboard", "No", "No", "No", "Yes", "Yes", "Yes", "L4 and above"],
        ["See all customers regardless of location", "No", "No", "No", "Yes", "Yes", "Yes", "L4+"],
        ["Admin: users, settings, recycle bin, imports", "No", "No", "No", "No", "No", "Yes",
         "Role must be exactly L6"],
    ],
    widths=[1.55, 0.32, 0.32, 0.32, 0.32, 0.32, 0.32, 2.6],
    font_size=8,
)

h2("Customer access sub-matrix")
para(
    "This is the layer most people miss. Access to a specific customer is decided before role "
    "level matters, and being a handler beats everything below L4.",
    size=9.5,
)
table(
    ["Role", "Is a handler of this customer", "Not a handler, location matches",
     "Not a handler, location does not match"],
    [
        ["L1", "FULL", "NONE", "NONE"],
        ["L2", "FULL", "NAME", "NONE"],
        ["L3", "FULL", "FULL", "NAME"],
        ["L4", "FULL", "FULL", "FULL"],
        ["L5", "FULL (via role, not handler)", "FULL", "FULL"],
        ["L6", "FULL (via role, not handler)", "FULL", "FULL"],
    ],
    widths=[0.7, 2.1, 1.9, 1.9],
)
para(
    "NAME access exposes only the customer's name, location, type, priority and handler list — "
    "never contacts, cases or quotations. NONE means the customer does not appear in that "
    "user's searches or lists at all. Note that the handler check runs before any role check "
    "and has no minimum level, so an L1 who is stored as a handler gets FULL access to that "
    "customer.",
    size=9.5,
)

h2("Case visibility")
para("A case is visible to a user if any one of these is true:")
bullet("They are L4 or above.")
bullet("Their email is in the case's stored owner list.")
bullet("They are the case's current assignee.")
bullet("They have FULL access to the case's customer.")
para(
    "Below L4 this is entirely role-independent. An L1 who is a case owner can change that "
    "case's stage, set its outcome and reassign its ticket — none of those three actions adds a "
    "role check on top of visibility. This is called out again in the limitations section "
    "because it is the most likely thing to surprise someone reading the matrix.",
    size=9.5,
)

# ======================================================================================
# P3
# ======================================================================================
point(3, "Database security check for unauthorised access")

labelled("What was asked.", "\"Check security of database for unauthorised access.\"")

labelled(
    "What we found.",
    "No exploitable path to unauthorised data was found. The full review is docs/security-audit.md; "
    "every finding there carries a file and line reference and a concrete exploit scenario, or an "
    "explicit statement of why it is not exploitable. Findings that turned out not to be real were "
    "written up as such rather than padded into the report.",
)

h2("Findings")
table(
    ["Severity", "Finding"],
    [
        [
            "Medium",
            "Anyone who can complete a Google sign-in with an @automationsystems.org address is "
            "automatically given a live, active CRM account on their first request, with no admin "
            "approval. This is the single most important item. It is narrower than it sounds: a "
            "freshly created account is L1 with no locations, and the access rules correctly give "
            "such a user NONE access to every customer — they see nothing until an admin raises "
            "their role, gives them a location, or adds them as a handler. The real risk is "
            "account hygiene, not data exposure: a departed employee whose Google account is still "
            "alive, or a shared mailbox, becomes a standing CRM identity that nobody was told "
            "about. Nothing in the code notifies an admin that this happened.",
        ],
        [
            "Informational (by design, stated plainly)",
            "The database's own row-level security gives zero protection against the application "
            "server. The app connects to Postgres with a full-access pooled role, so every "
            "row-level policy is bypassed. That is the intended architecture — those policies "
            "exist to make the public Supabase key that ships to every browser useless — but it "
            "means the entire authorisation boundary is application code, not the database. A bug "
            "in one permission check is the only thing between a signed-in user and any row.",
        ],
        [
            "Low",
            "/api/download/... is not listed in the middleware's protected-path list. Checked "
            "directly: the route re-derives the caller's identity itself and re-checks access to "
            "the specific quote, so it is not a data leak. The only consequence is that an "
            "unauthenticated browser hit gets a raw JSON error instead of a login redirect.",
        ],
        [
            "Low",
            "One Google Drive API query string is built by string interpolation. Not exploitable: "
            "the value interpolated comes from a settings row written only by the L6-gated Drive "
            "setup flow, never from user input. Flagged as a pattern to watch if a future feature "
            "ever threads a user-supplied folder ID through it.",
        ],
        [
            "Informational",
            "Row-level security is complete across all 14 tables (deny-all for both public "
            "database roles, grants and sequences revoked), and none of the later migrations "
            "introduces an uncovered table.",
        ],
        [
            "Informational",
            "Identity is always derived on the server per request and is never taken from what the "
            "browser sends. Every place that accepts an email in a request payload was checked: in "
            "each case it names the target of the action, and the caller is re-authorised "
            "separately.",
        ],
        [
            "Informational",
            "No SQL injection surface. Every database call in the entire server tree uses "
            "parameterised tagged templates; there are zero uses of raw or concatenated SQL.",
        ],
        [
            "Informational",
            "Secrets: the working secrets file is git-ignored and untracked. The Google Drive "
            "setup routes are L6-gated, refuse to run once configured, and never store or log the "
            "refresh token. Every admin operation calls the admin check as its first statement.",
        ],
    ],
    widths=[1.25, 5.3],
)

h2("What could not be checked from the code alone")
para(
    "This was a code review. The following need someone with dashboard access to confirm, and "
    "are worth ten minutes of somebody's time:"
)
bullet(
    "Whether row-level security is actually enabled on the live production database, as opposed "
    "to correctly written in the migration file. A policy dropped by hand during debugging and "
    "never restored would not be visible from the repository."
)
bullet(
    "Whether the Google Cloud OAuth consent screen is genuinely set to \"Internal\" audience. If "
    "it is, Google itself blocks non-company accounts before the CRM is ever involved, which "
    "meaningfully narrows the Medium finding above."
)
bullet(
    "Whether the Supabase service-role key listed in the deployment environment is still needed. "
    "No code path in the repository uses it. An unused key is housekeeping, not a vulnerability, "
    "but it should either be used or removed."
)

h2("Recommendation")
para(
    "One change is worth making, and it is a product decision rather than a bug fix: either "
    "notify an admin when a new account auto-provisions, or create such accounts inactive and "
    "require an L6 to switch them on. The first costs nothing and preserves first-login "
    "convenience; the second trades a little friction for a control that is enforced by code "
    "rather than by remembering to check the Users screen. Neither was implemented here, because "
    "it was not part of what you asked for."
)

# ======================================================================================
# P4
# ======================================================================================
point(4, "Scalability and storage estimate")

labelled(
    "What was asked.",
    "\"Check & discuss scalability of the system for more users with an estimate of storage.\"",
)

labelled(
    "What we found.",
    "The database schema is fine. The query pattern is not, and that — not storage — is what will "
    "bite first. The customer list and the case list both fetch the entire table with no WHERE "
    "clause and then filter in JavaScript. On top of that, the case list and the dashboard each "
    "fetch the customer record for every single case, one query per case. The row limits that do "
    "exist (400 customers, 300 cases) are applied after the whole table has already been pulled "
    "into memory, so they reduce the response size, not the cost.",
)
para(
    "A side effect worth knowing: 14 database indexes were added in earlier work, and most of "
    "them cannot be used by any query the app currently issues, because a query with no WHERE "
    "clause has nothing to look up. That work is not wasted — it is premature. Those indexes "
    "start paying for themselves the day the queries are rewritten to filter in SQL. Until then "
    "they are pure write overhead."
)

h2("Storage estimate")
para(
    "Row sizes below are computed from the actual column types in the schema. The usage rates "
    "are stated assumptions for a B2B industrial-automation sales team, not measured production "
    "numbers — this analysis has no access to the live database's row counts.",
    size=9.5,
)
table(
    ["Assumption (per active user per year)", "Value"],
    [
        ["New customers", "15"],
        ["New cases", "40"],
        ["Contacts per customer", "2"],
        ["Quotations per case", "0.6"],
        ["Uploaded (not Drive-saved) quotes, as a share of quotations", "30%"],
        ["Average uploaded file size", "1 MB"],
        ["Activity-log rows per case / per customer", "6 / 3"],
        ["Index storage overhead applied on top of row storage", "+60%"],
    ],
    widths=[4.6, 1.9],
)

para("That works out to roughly 7.3 MB per user per year, of which 6.9 MB is uploaded files.")

table(
    ["Total database size", "1 year", "3 years", "5 years"],
    [
        ["20 users", "~147 MB", "~440 MB", "~734 MB"],
        ["50 users", "~367 MB", "~1.10 GB", "~1.83 GB"],
        ["100 users", "~734 MB", "~2.20 GB", "~3.67 GB"],
    ],
    widths=[1.8, 1.5, 1.5, 1.5],
)

para(
    "The single most important number here: uploaded quotation files are stored as raw bytes "
    "inside the Postgres row, and they outweigh every other table combined by roughly 35 to 1. "
    "Everything else in the database is noise by comparison. If uploaded files were pushed to "
    "Google Drive and the copy in Postgres cleared afterwards — the Drive save feature already "
    "exists, it just does not delete the inline copy — total storage drops by about two orders "
    "of magnitude, to roughly 150-750 KB per user per year. That one behavioural choice is a "
    "bigger lever on storage than user count and years combined."
)

h2("Where it breaks, and in what order")
table(
    ["Endpoint", "Fine until roughly", "What gives way first"],
    [
        ["Recent activity feed (dashboard)",
         "Already a problem today",
         "Up to 250 sequential database round trips for an L1-L3 user. Fixed in this batch — see P12"],
        ["Cases tab and dashboard",
         "~2,000-5,000 cases",
         "One database round trip per case. Starts costing noticeable time once any team has a few hundred visible cases, well before row counts get large"],
        ["Customers tab",
         "~5,000-10,000 customers",
         "Full-table fetch then JavaScript filtering on every page load, uncached, per user. Roughly a 2-4 year horizon at the growth rates above"],
        ["Uploaded quotation storage",
         "Billing decision, not a break",
         "Grows with upload behaviour rather than user count. Dominant byte cost but never causes a failure"],
    ],
    widths=[1.5, 1.4, 3.6],
)

h2("Recommended fixes, in priority order")
numbered(
    "Batch the recent-activity customer lookups. Highest impact for the least effort — this was "
    "a live latency bug, not a future one, and it is the item that was actually fixed in this "
    "batch (P12)."
)
numbered(
    "Replace the one-query-per-case pattern in the cases list and the dashboard with a single "
    "batched lookup, or push the join into SQL. Medium effort. This removes the dominant latency "
    "driver on the two pages everyone loads most."
)
numbered(
    "Give the customer and case list queries real WHERE clauses and pagination, pushing the "
    "JavaScript filters down into SQL. Medium-to-large effort, and it needs care around the "
    "visibility logic, which currently depends on having the handler table loaded in memory. "
    "This is the 3-5 year fix and the one that makes the existing indexes earn their keep."
)
numbered(
    "Memoise the handler and user table reads per request. Small effort, low impact at current "
    "scale, but cheap to do while touching the same call sites for items 2 and 3."
)
numbered(
    "Decide a retention policy for uploaded quotation bytes — clear the inline copy once a file "
    "has a Drive link. Largest single lever on storage cost; no effect on latency."
)
para(
    "Items 2 through 5 were not implemented in this batch. They were not part of your review "
    "document, and item 3 in particular is a large enough change that it deserves to be its own "
    "piece of work rather than a rider on this one.",
    size=9.5,
)

h2("Plan limits worth confirming")
para(
    "The repository cannot see which Supabase or Vercel plan is active. Given the projection "
    "above, three things are worth confirming with whoever holds the billing access: the "
    "Supabase database size cap on the current tier (the 3-5 year projection at 50-100 users "
    "approaches the level where free and entry tiers matter), the transaction pooler's "
    "connection ceiling, and the Supabase monthly egress allowance — full-table reads on every "
    "page load are the most bandwidth-hungry pattern in the app today. None of these is an "
    "alarm; they are line items for a ten-minute check."
)

# ======================================================================================
# P5
# ======================================================================================
point(5, "Remove the redundant dashboard buttons")

labelled(
    "What was asked.",
    "\"Remove buttons: New customer, Customers, Cases as these are already covered in the top bar.\"",
)
labelled(
    "What we found.",
    "Correct on all three. \"+ New customer\" sat on the L4 dashboard and \"Customers\" / "
    "\"Cases\" sat on the L5/L6 dashboard. All three duplicated the top navigation exactly.",
)
labelled("What changed.", "All three buttons removed from the dashboard views.")
labelled(
    "Current behaviour.",
    "Neither the L4 nor the L5/L6 dashboard shows any of the three buttons. The top-bar routes "
    "to Customers and Cases are unchanged and still work. An end-to-end test asserts both halves "
    "of that — the buttons are gone and the navigation still routes.",
)

# ======================================================================================
# P6
# ======================================================================================
point(6, "Fix the oversized checkbox; split it into two")

labelled(
    "What was asked.",
    "\"UI fix of check mark, also need 2 check boxes — owned by me & assigned to me.\"",
)

labelled(
    "What we found.",
    "The oversized blue slab was not a checkbox-specific style at all. Two global rules were "
    "landing on it: a general rule that stretched every input to full width with a 40px minimum "
    "height and padding, and a second rule inside the filter bar that forced a 140px minimum "
    "width on every input there. The checkbox was inheriting both. This matters because it "
    "explains why the obvious fix does not work: adding an inline width to the element does "
    "nothing about the minimum-width and minimum-height rules, and the filter-bar rule has the "
    "same CSS specificity as a plain attribute selector, so the override has to both out-specify "
    "it and come later in the stylesheet.",
)

labelled(
    "What changed.",
    "A single corrective rule that out-specifies and follows both offenders, pinning checkboxes "
    "and radio buttons to 16px square with no padding and no minimum dimensions. It is written "
    "to cover every checkbox in the application, not just this one — the customer grid's "
    "row-select boxes had the same problem. Separately, the single \"Mine only\" checkbox became "
    "two independent ones, \"Owned by me\" and \"Assigned to me\", with matching filter flags "
    "added to the server's case-list call.",
)

labelled(
    "Current behaviour.",
    "The two boxes combine with OR, as agreed. Ticking \"Owned by me\" alone shows cases you own; "
    "\"Assigned to me\" alone shows cases where you hold the ticket; both ticked shows either; "
    "neither ticked shows everything you can see. The old \"mine\" flag is still accepted by the "
    "server and treated as \"owned\", so a browser with an older copy of the page cached keeps "
    "working rather than silently losing its filter.",
)

# ======================================================================================
# P7
# ======================================================================================
point(7, "Location mandatory, moved next to Name, and renamed from Tags")

labelled(
    "What was asked.",
    "\"At least one tag should be mandatory for a customer, move this field up next to name "
    "while creating customer; also rename tags -> location everywhere.\"",
)

labelled(
    "What we found.",
    "The field was optional on both create and update, sat well down the form, and was labelled "
    "\"Tags\" in seven different places. It is also the field that drives access control — who "
    "can see which customer — which is a good reason for it to be mandatory and a good reason "
    "for it to be named something a salesperson recognises.",
)

labelled(
    "What changed.",
    "Creating a customer now requires at least one location, and updating one cannot leave it "
    "empty. The field moved directly beneath Customer name in the create form. Every user-facing "
    "label now reads \"Location\": the form, the grid header, the filter dropdown, the detail "
    "view, the import instructions and the admin screen.",
)

para(
    "One deliberate decision here, taken with the owner: the rename is labels only. The database "
    "column is still called tags and the API field name is unchanged. Renaming a column that "
    "drives access control, on a live database, to achieve a cosmetic change, is risk with no "
    "return. Nobody using the CRM will ever see the word \"tags\" again; anyone reading the "
    "database will."
)

para(
    "Existing customers with no location would have become un-savable the moment the rule turned "
    "on, so migration 0007 fills them in with the placeholder TO BE FILLED. That value is "
    "recognised as valid — so a backfilled record survives being saved again without being "
    "silently emptied — but it is deliberately excluded from the picker, so nobody can choose it "
    "on purpose. It exists to be visible and to be replaced."
)

labelled(
    "Current behaviour.",
    "Location is required, sits second on the form under Name, and is called Location everywhere "
    "a user can see it. No existing record was blocked or emptied. Any customer showing "
    "TO BE FILLED is a record that had no location before and needs a real one.",
)

# ======================================================================================
# P8
# ======================================================================================
point(8, "SEI as an optional dropdown multi-select")

labelled(
    "What was asked.",
    "\"Add new multi select field: 'SEI'. This will be for Schneider person name. In UI while "
    "creating [a customer], unlike pill shaped element of tags, this should be a drop down multi "
    "select list; and not mandatory.\"",
)

labelled(
    "What we found.",
    "SEI already existed as a free-text field, edited as a plain text box in the customer grid. "
    "Free text meant the same person's name could be spelled three ways and no filter would ever "
    "work reliably.",
)

labelled(
    "What changed.",
    "Migration 0008 converts the column from single text to a list. Existing values are split on "
    "pipes and commas and preserved, never dropped — the migration verifies this row by row and "
    "aborts if any row disagrees with the expected result. The selectable names live in the "
    "settings table, editable by an L6 under Admin, and the application reads that list live "
    "rather than from a hardcoded constant. The field is rendered as a dropdown multi-select, "
    "explicitly not the pill-style picker used by Location, as you asked.",
)

para(
    "The list ships empty. No Schneider names were invented. Until an L6 populates it under "
    "Admin, the field shows \"No SEI names configured yet\" and cannot be set — this is the first "
    "thing to do after deployment, and it is repeated in the limitations section so it does not "
    "get lost."
)

labelled(
    "Current behaviour.",
    "SEI is optional; a customer may have zero, one, or many names. A name that is not on the "
    "admin list is rejected with a message pointing at where to add it, rather than being "
    "silently dropped — so a typo cannot quietly lose data.",
)

# ======================================================================================
# P9
# ======================================================================================
point(9, "Direct as a special account")

labelled(
    "What was asked.",
    "No Remove button on Direct, because it clears itself the moment a real handler is mapped. "
    "Direct should appear in the user list as a special account with all locations and no login "
    "credentials, so that L4 and above can view the Direct dashboard from their home screen "
    "using the \"view someone else's dashboard\" option. And Direct should only be set when no "
    "real handler exists.",
)

labelled(
    "What we found.",
    "The invariant you described was already correct — Direct is set only when there is no real "
    "handler, and adding a real handler clears it automatically. But the customer screen still "
    "offered a Remove button next to Direct, and the only thing that button could achieve was "
    "leaving the customer with no handler at all. Direct also did not appear in the user list or "
    "the dashboard picker. And the dashboard for Direct could never have worked — see the "
    "separate bugs section.",
)

labelled(
    "What changed.",
    "The Remove button is withheld when the handler is Direct, and the fake email address is not "
    "shown either; instead the row reads \"Unassigned — no account handler yet. Adding one "
    "replaces this automatically.\" Direct is now synthesised as a virtual user: name Direct, all "
    "locations, explicitly flagged as having no login. It appears in the user list and in the "
    "dashboard picker for L4 and above.",
)

para(
    "Direct deliberately has no row in the users table and never will. That table has a "
    "constraint requiring a real company email address, and Direct is not a person. It is stored "
    "only as a handler value and synthesised into user lists when they are read. This is what "
    "makes \"no login credentials\" structural rather than a setting somebody could flip."
)

labelled(
    "Current behaviour.",
    "L4 and above can select Direct in the dashboard picker and see the customers and cases it "
    "holds; below L4 the request is refused. Direct can never be assigned a ticket and can never "
    "be a case owner — it is excluded from every list of people you can assign work to. It is a "
    "subject you can look at, never a target you can hand work to.",
)

# ======================================================================================
# P10
# ======================================================================================
point(10, "BUG — a case on a Direct-handled customer mislabelled its creator")

labelled(
    "What was asked.",
    "You created a customer from an L6 account (himanshuneb), assigned it to nobody, then "
    "created a case on it. The case reported that the account handler was Himanshu Neb. Your "
    "note: on a Direct-handled account the default case owner should be the case creator, and it "
    "should not say he is the account handler; also make sure a case always has at least one "
    "owner.",
)

labelled(
    "What we found.",
    "Real bug, and the root cause was structural rather than a wrong label. One function was "
    "answering two different questions at once. It derived case owners from the customer's "
    "handler list; when the only handler was Direct it filtered Direct out, found nobody, and "
    "fell back to the case's creator. Fine so far — the creator is the right owner. The problem "
    "is that the user interface asked the same function \"is this person an account handler?\" in "
    "order to decide whether they could be removed from the case. Because the answer came back "
    "through the same channel, the creator was marked non-removable and labelled \"account "
    "handler — owner of every case on the account\". So the screenshot was not only saying the "
    "wrong thing; the person it named could not be taken off the case either.",
)

labelled(
    "What changed.",
    "The two concepts are now separate. Every owner on a case carries an explicit source: "
    "handler, creator, or manual. The interface labels each one accordingly — \"(account handler "
    "— owner of every case on the account)\", \"(created this case)\", or \"(added to this "
    "case)\" — and removability is decided from the source, not from a shared fallback. An owner "
    "whose source is creator or manual can be removed; one whose source is handler cannot, "
    "because they own the case by virtue of owning the account.",
)

labelled(
    "Current behaviour.",
    "A case created on a Direct-handled customer is owned by whoever created it, labelled "
    "\"created this case\", and that person can be removed from the case if someone else is "
    "added first. The at-least-one-owner rule is enforced on the server: removing the last "
    "remaining owner is refused outright. The exact scenario from your screenshot has a unit "
    "test reproducing it — L6 creates a customer, gets Direct as handler, creates a case — "
    "asserting that the owner comes back labelled creator and not handler.",
)

# ======================================================================================
# P11
# ======================================================================================
point(11, "Handler changes must propagate to case owners")

labelled(
    "What was asked.",
    "\"Since case owner is linked with customer handler ... check the behaviour when a customer "
    "handler is added/removed after cases are created. Ideally, all the cases for that customer "
    "should get updated to add [the] new handler if added. But it should not remove [the] case "
    "owner if [the] customer handler is removed. And this update process should be done only for "
    "the active case, not the ones that have been closed in the past.\"",
)

labelled(
    "What we found.",
    "The existing architecture could not express any of those three rules, and this is the single "
    "most consequential thing in this batch. Case ownership was not stored anywhere. It was "
    "recomputed on every read from the customer's current handler list. That has three "
    "consequences, all of them wrong by your specification: adding a handler retroactively made "
    "them an owner of every case on the account, including ones closed years ago; removing a "
    "handler silently erased them from every case they had ever owned, closed cases included, "
    "with no record that they were ever there; and \"active cases only\" was meaningless, because "
    "there was no update event to scope — the answer simply changed the next time anyone looked.",
)

labelled(
    "What changed.",
    "Ownership is now stored on the case itself. When a case is created, its owners are seeded "
    "with the customer's real handlers, or with the creator when the only handler is Direct. "
    "Adding a handler appends them to the owner set of that customer's cases that are still open. "
    "Removing a handler does not touch any case. Every read path — the case screen, the "
    "dashboard, the case list — reads the stored set.",
)

para(
    "The old derivation was not deprecated, it was deleted. The function that performed it is "
    "gone and the function that replaced it no longer accepts the argument that would let a "
    "caller ask for the old behaviour. Making the wrong thing structurally impossible is more "
    "durable than a comment asking people not to do it."
)

para(
    "Migration 0005 seeds the stored owner set for every existing case with exactly what the old "
    "derivation would have produced, so nothing visibly changes on the day it ships. Its "
    "verification is per case, not in aggregate: the owner set computed before the migration must "
    "equal the stored set after it, for every single case, or the whole thing rolls back. This is "
    "the highest-risk change in the batch and it is checked accordingly."
)

labelled(
    "Current behaviour.",
    "Add a handler to a customer and they become an owner of that customer's open cases. Closed "
    "cases are left byte-for-byte identical. Remove a handler and every case they own is "
    "untouched — they keep their ownership and their access to those cases. A case always has at "
    "least one owner. Two details worth knowing: a case on Hold counts as closed for this "
    "purpose, so a handler added while a case is on Hold will not pick it up even after it "
    "reopens; and when someone stops being a handler, their existing case ownership is relabelled "
    "from \"account handler\" to \"created this case\" or \"added to this case\", which also "
    "means it becomes removable by hand.",
)

# ======================================================================================
# P12
# ======================================================================================
point(12, "Dashboard issued up to 250 sequential database queries")

para(
    "This point was not in your review document. It was found during the scalability analysis "
    "you asked for in P4, and it was fixed in this batch with the owner's approval because it "
    "was cheap to fix and was affecting people today.",
    italic=True,
)

labelled(
    "What we found.",
    "The recent-activity panel on the dashboard fetches the last 250 activity entries and then "
    "filters out the ones the viewer is not allowed to see. To decide that, it needed the "
    "customer record for each entry — and it fetched them one at a time, inside the filter loop, "
    "waiting for each before starting the next. Not in parallel: one after another. For an L1, L2 "
    "or L3 user looking at a busy team feed, that is up to 250 sequential round trips to a "
    "database in a different country. At realistic cross-region latency that is several seconds "
    "of dead time on the page everyone sees first after logging in. It also gets worse as the "
    "company grows, because more colleagues means more activity entries that a lower-level "
    "viewer cannot see directly and therefore has to check.",
)

labelled(
    "What changed.",
    "Every customer record the loop could possibly need is knowable before the loop starts. They "
    "are now collected up front, de-duplicated, and fetched in one batched query. The filtering "
    "logic itself is unchanged — the same people see the same entries.",
)

labelled(
    "Current behaviour.",
    "The recent-activity panel issues one query instead of up to 250. There is a second, quieter "
    "improvement: the old code also returned an empty activity list in some cases where it should "
    "have returned entries, so some users were seeing a blank panel that now populates correctly.",
)

# ======================================================================================
# Incidental bugs
# ======================================================================================
h1("Three pre-existing bugs found along the way")

para(
    "None of these were in your document and none were symptoms you had reported. They were "
    "found while working on the points above, and all three are fixed. They are listed because "
    "you should know they existed, and because two of them had been quietly wrong for some time."
)

h2("1. Every bulk import silently made an L6 admin the account handler")

para(
    "The bulk customer importer took handler emails from the spreadsheet, and when a row had "
    "none, it defaulted to whoever was running the import. Only an L6 admin can run an import. "
    "So every customer imported without an explicit handler in the sheet was assigned to an L6 "
    "admin as its account handler — precisely the thing P1 says must never happen — and it "
    "happened at bulk scale, invisibly, every time. Sheet-supplied emails that resolved to an L5 "
    "or L6 user had the same problem."
)
para(
    "Fixed on both paths: handler emails from the sheet that resolve to L5 or L6 users are "
    "filtered out, and the default when a row names nobody is now Direct rather than the "
    "importing admin. Migration 0006 cleans up any such rows already in the database."
)

h2("2. The Direct dashboard could never have worked")

para(
    "When the dashboard is asked to show a specific person, it normalises the email it is given "
    "by appending the company domain if it is missing. Direct is stored as the literal string "
    "\"direct\", which has no @ in it. So a request for the Direct dashboard was normalised into "
    "a request for direct@automationsystems.org — an address that does not exist and never has. "
    "The Direct dashboard could not have returned anything meaningful."
)
para(
    "This was found while implementing P9, which is what made it visible: your request to expose "
    "Direct in the dashboard picker would have shipped a picker entry that led to an empty page. "
    "The fix is a one-line ordering change — Direct is recognised before the domain is appended, "
    "not after."
)

h2("3. Migration 0008 had to convert the recycle bin too")

para(
    "The SEI column change (P8) converts customers.sei from text to a list. The recycle bin keeps "
    "its own copy of every column of a deleted customer, and restoring a customer copies those "
    "columns straight back into the customers table. Converting only the live table would have "
    "left the recycle bin holding the old text format, and every single customer restore would "
    "have failed on a type mismatch — an error nobody would have seen until the first time "
    "somebody tried to undelete a customer, possibly weeks later."
)
para(
    "Migration 0008 converts both columns, using the same splitting rule for each. This is the "
    "kind of thing that is easy to miss and expensive to find later, so it is worth flagging "
    "that it was caught before deployment rather than after."
)

# ======================================================================================
# Deployment
# ======================================================================================
h1("Deployment")

para(
    "Four database migrations were written for this batch. None of them has been applied to any "
    "database — not production, not staging, not a local copy. They are written, unit-tested "
    "against their reference implementations, and each one verifies its own result in SQL before "
    "committing, but the deployment step itself has not been taken.",
    bold=False,
)

h2("Migration order")

para(
    "They must run in numeric order. This is not a convention — 0005 before 0006 is a hard "
    "correctness requirement, explained below the table.",
)

table(
    ["#", "Migration", "What it does", "Self-check before it commits"],
    [
        ["0005", "Materialise case owners",
         "Writes each existing case's owner set onto the case row, using exactly what the old "
         "read-time derivation would have produced.",
         "Recomputes the old derivation for every case and compares it to the stored result, case "
         "by case. Any single mismatch aborts the transaction."],
        ["0006", "Remove L5/L6 handlers",
         "Deletes handler rows belonging to L5 and L6 users.",
         "Asserts there were such rows before and exactly zero after, and asserts that the number "
         "of cases holding at least one owner is unchanged. This file never writes to the cases "
         "table at all, and a test asserts that it doesn't."],
        ["0007", "Backfill customer locations",
         "Fills empty location lists with the placeholder TO BE FILLED.",
         "Asserts there were empty ones before and exactly zero after, and that the total customer "
         "count is unchanged — nothing created, nothing deleted."],
        ["0008", "SEI multi-select",
         "Converts the SEI column from text to a list, on both the customers table and the recycle "
         "bin, splitting existing values on pipes and commas. Seeds an empty SEI name list in "
         "settings.",
         "Snapshots every pre-conversion value, computes the expected result for each, and compares "
         "row by row after the conversion."],
    ],
    widths=[0.45, 1.35, 2.45, 2.3],
    font_size=8,
)

h3("Why 0005 must run before 0006")
para(
    "Under the old model, case ownership was derived from the handler table at read time. "
    "Deleting a handler row therefore did not just remove a handler — it silently erased that "
    "person from every case they owned, including cases closed years ago, with nothing left "
    "recording that they had ever owned them. That is irreversible, and it is exactly the "
    "behaviour P11 asks us to prevent."
)
para(
    "Once 0005 has stored ownership on each case, deleting the handler row is inert: the case "
    "keeps the owners it already had. Running 0006 first would destroy historical ownership data "
    "with no way to get it back. Running them in order is safe. There is no recovery path if this "
    "is done in the wrong order, which is why it is stated this plainly."
)

h2("Recommended deployment sequence")
numbered("Take a full database backup and confirm you can actually restore it. Do this first.")
numbered(
    "Apply 0005, 0006, 0007, 0008 in that order to a staging or branch database that has real "
    "production data in it, not an empty one. The self-checks in these migrations only prove "
    "anything against real rows."
)
numbered(
    "Confirm on staging: existing cases show the same owners as before; no L5 or L6 appears as "
    "an account handler anywhere; customers that had no location now show TO BE FILLED; existing "
    "SEI values survived the conversion; and a customer restore from the recycle bin still works."
)
numbered("Deploy the application code.")
numbered(
    "Apply the same four migrations to production, in the same order. Each aborts its own "
    "transaction on any mismatch, so a failure leaves the database as it was rather than half "
    "converted."
)
numbered(
    "Immediately after: have an L6 populate the SEI names list under Admin > Settings. Until "
    "that is done the field cannot be used."
)
numbered(
    "Then: review any customer showing TO BE FILLED as its location and give it a real one."
)

para(
    "The application code and the migrations should go out close together. The code expects the "
    "new column types; the migrations do not depend on the new code. Applying migrations first, "
    "then deploying, is the safer order of the two.",
    size=9.5,
)

# ======================================================================================
# Limitations
# ======================================================================================
h1("Known limitations and follow-ups")

para(
    "These are real and outstanding. Some are deliberate trade-offs, some are things that were "
    "out of scope, and one is a piece of misleading text that is now technically wrong. They are "
    "listed here rather than buried because you will find them eventually and it is better that "
    "you hear them from us."
)

h2("1. The SEI list ships empty")
para(
    "No Schneider names were invented, so the dropdown has nothing in it until an L6 populates it "
    "under Admin > Settings. Until then the field displays \"No SEI names configured yet\" and "
    "cannot be set. This is the first post-deployment action and it takes about two minutes, but "
    "if it is skipped the feature will look broken."
)

h2("2. The case-owners screen still gives advice that no longer works")
para(
    "When you try to remove an owner who is there because they are the account handler, the "
    "system refuses and says to \"remove them as a handler on the customer instead\". That text "
    "predates this batch. It used to be true. Now that P11 is implemented as you specified — "
    "removing a handler must not strip an existing case owner — following that instruction will "
    "remove them as a handler and leave their case ownership exactly where it was. The behaviour "
    "is correct and is what you asked for; the sentence is stale."
)
para(
    "It was left alone deliberately rather than reworded on the fly, because the right wording "
    "depends on what you actually want to happen when someone genuinely needs a handler-sourced "
    "owner taken off a case. Right now there is no way to do that at all. Worth a short "
    "conversation, then a small change."
)

h2("3. Several case and quote actions have no role check at all")
para(
    "Changing a case's stage, setting its outcome, reassigning its ticket, setting a quote's "
    "status, generating a PDF and saving to Drive are gated only by whether you can see the "
    "record, not by your level. An L1 who is a case owner can do all of those. This is "
    "long-standing behaviour, not something introduced here, and it may well be what you want for "
    "a small trusted team — but it will surprise anyone who reads the role matrix and assumes "
    "\"L-level equals permission\"."
)
para(
    "There is a related cosmetic inconsistency: the case screen returns a flag suggesting that "
    "ticket assignment requires L2 or above. The server does not enforce that. The flag only "
    "affects what buttons are drawn."
)

h2("4. Authorisation is per-service, with no central chokepoint")
para(
    "Authentication is centralised — every request re-derives who you are from the session, "
    "enforces the company domain, and rejects inactive users, at one place. Authorisation is not. "
    "Every \"who can do what\" rule lives inside an individual service function. Every action in "
    "the matrix was traced individually and every one of them is gated, so this is not a present "
    "defect. It is a maintenance risk: a new endpoint added later without a permission check "
    "would be reachable by any signed-in company user, and nothing structural would catch it."
)

h2("5. The database is in Tokyo and the application is in Mumbai")
para(
    "The Supabase database runs in ap-northeast-1 (Tokyo) while the Vercel functions run in bom1 "
    "(Mumbai). Every single database round trip pays that distance — realistically tens of "
    "milliseconds each — and the app makes many round trips per page. This is why the sequential "
    "query bug in P12 was as bad as it was, and it is a standing tax on every other query too. "
    "Moving the database to a region near Mumbai would be the single largest latency improvement "
    "available, larger than any query rewrite. It is also a migration with real risk and downtime, "
    "so it is a decision rather than a task."
)

h2("6. Query patterns from P4 that were not fixed")
para(
    "Only the recent-activity bug was fixed. The one-query-per-case pattern on the cases list and "
    "the dashboard, the unfiltered full-table reads behind the customers and cases lists, and the "
    "inline storage of uploaded quotation files all remain as described in P4. None is urgent at "
    "twenty users; the first of them starts to matter once any team accumulates a few hundred "
    "visible cases."
)

h2("7. Storage numbers are modelled, not measured")
para(
    "The projections in P4 are computed from the real column types in the schema, but the usage "
    "rates behind them are stated assumptions. Real usage could plausibly run two to five times "
    "higher or lower, driven almost entirely by how often people upload vendor PDFs instead of "
    "generating quotes in the app. Treat the numbers as a shape, not a forecast."
)

h2("8. Items that could not be verified without dashboard access")
para(
    "Whether row-level security is actually live on the production database, whether the Google "
    "OAuth consent screen is genuinely restricted to company accounts, and which Supabase and "
    "Vercel plan tiers are active. All three are listed with their consequences in P3 and P4."
)

h2("9. Won order value is counted in full on every owner's dashboard")
para(
    "This is pre-existing behaviour and was not part of the review, but this batch makes it more "
    "visible, so it should be said plainly. When a case is won, its full order value is added to "
    "the dashboard total of every owner - it is not divided between them. A case worth 30 lakh "
    "with three owners contributes 30 lakh to all three dashboards, and the same order is counted "
    "three times if you add those dashboards together."
)
para(
    "That was a deliberate decision taken earlier in the project (recorded as \"its fine dont "
    "split\"), and nothing here has changed it. The reason it matters now is P11: adding an "
    "account handler makes that person an owner of the customer's open cases, so multi-owner "
    "cases will become more common than they were. Individual dashboards remain correct as a "
    "measure of what each person is involved in. They simply cannot be summed to get a company "
    "total. If a combined figure is ever needed, that needs its own report rather than an "
    "addition of these ones."
)

h2("10. Admin settings are read back for SEI names only")
para(
    "The Admin settings screen lets an L6 edit locations, types, priorities, product categories "
    "and sources, and those edits do save to the database. But most of the application reads "
    "those lists from values compiled into the code rather than from the database, so with one "
    "exception the edits have no visible effect until a developer changes the code and redeploys."
)
para(
    "The exception is the new SEI names list from P8, which was deliberately wired to read live "
    "from the database - an admin-managed list that the application never read back would have "
    "been broken by design. The rest was left alone on purpose: changing it touches customer and "
    "case validation across the whole system, which is not a change worth bundling into this "
    "batch. It is a known gap with a clear fix, not a defect introduced here. Until it is done, "
    "treat Admin > Settings as functional for SEI names and cosmetic for everything else."
)

# ======================================================================================
# Closing
# ======================================================================================
h1("Summary of what happens next")

table(
    ["Owner", "Action"],
    [
        ["You / project owner", "Confirm the deployment window and that a restorable backup exists."],
        ["Deployment", "Apply migrations 0005-0008 in order to staging with real data, verify, then production."],
        ["L6 admin", "Populate the SEI names list under Admin > Settings, immediately after deployment."],
        ["L6 admin", "Review customers showing TO BE FILLED and set a real location."],
        ["You / project owner",
         "Decide on the case-owners wording (limitation 2), and whether auto-provisioned accounts "
         "should start inactive (P3 recommendation)."],
        ["Follow-up work",
         "The remaining P4 query fixes, and the question of the database region, if and when they "
         "become worth doing."],
    ],
    widths=[1.5, 5.0],
)

para(
    "Supporting documents in the repository: docs/role-matrix.md (full permission matrix with "
    "line-level citations), docs/security-audit.md (full security review), "
    "docs/scalability-and-storage.md (full storage and query analysis).",
    size=9.5,
)

doc.save(OUT)
print(f"Wrote {OUT}")
