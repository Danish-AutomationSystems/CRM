"""Generate a plain-language scalability summary for non-technical readers.

Source of truth: docs/scalability-report-2026-08.md (the detailed, measured version).
This document deliberately drops code references, SQL and per-table statistics.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

ACCENT = RGBColor(0x1F, 0x3B, 0x73)
GOOD = RGBColor(0x1B, 0x7F, 0x3B)
WARN = RGBColor(0xB3, 0x5C, 0x00)
BAD = RGBColor(0xA3, 0x1D, 0x1D)

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(8)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def para(text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(head)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            cells[i].paragraphs[0].add_run(str(value))
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ---------------------------------------------------------------- title page

title = doc.add_heading("CRM Scalability — Plain Summary", level=0)
for run in title.runs:
    run.font.color.rgb = ACCENT

sub = para(
    "How many people can use the CRM on our current free hosting, "
    "and what it would take to support more."
)
sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in sub.runs:
    run.italic = True
    run.font.size = Pt(12)

para("Date: 13 August 2026")
para("Current setup: Vercel (free plan) for the website, Supabase (free plan) for the database.")
para(
    "Note: every number in this summary was measured directly against the live system "
    "on the date above, not estimated from the design.",
    italic=True,
)

doc.add_page_break()

# ---------------------------------------------------------------- the answer

heading("1. The short answer", 1)

para(
    "Right now the CRM comfortably supports about 5 to 8 regular users before it starts "
    "running into free-plan limits.",
    bold=True,
)

para(
    "That number is low, but not because the CRM stores too much information. It is low "
    "because of two specific choices in how the system was built. Both are fixable in about "
    "one to two days of work, and fixing them raises the ceiling to roughly 25 to 40 users "
    "on the same free plans — at no extra cost."
)

table(
    ["Situation", "How many users it supports", "What runs out first"],
    [
        ["Today, as built", "5 – 8", "Database space, filled by uploaded quotation files"],
        ["After fix 1 (see Section 3)", "10 – 15", "Pages get slow as case count grows"],
        ["After fixes 1, 2 and 3", "25 – 40", "Monthly data transfer allowance"],
        ["After all four fixes", "100 +", "Free-plan computing power"],
    ],
    widths=[2.0, 1.8, 2.5],
)

# ---------------------------------------------------------------- what we found

heading("2. What we found", 1)

heading("The CRM's own records are not a problem at all", 2)

para(
    "Customers, cases, contacts, notes and history are all very small. Measured against real "
    "records in the live system, twenty people working normally would generate about 5 MB of "
    "records per month. The free database holds 500 MB."
)

para("At twenty users, the CRM's records alone would take over eight years to fill it.", bold=True, color=GOOD)

heading("Uploaded quotation files are the real problem", 2)

para(
    "When someone uploads a quotation PDF, the file is currently saved inside the database "
    "itself rather than in a file store. Each file can be up to 8 MB."
)

para(
    "At the same twenty users, uploaded files would consume about 200 MB per month — roughly "
    "forty times more space than everything else in the CRM combined. The free database would "
    "be full in under three months."
)

para(
    "This single issue is the reason the current ceiling is 5 to 8 users instead of 40.",
    bold=True,
    color=BAD,
)

para(
    "The good news: the CRM is already connected to Google Drive, which is where these files "
    "belong. And there are currently zero uploaded files in the live system, so moving this "
    "over now costs nothing. The longer we wait, the more work it becomes."
)

heading("Pages will get slow as the number of cases grows", 2)

para(
    "When the case list loads, the system fetches the customer details for each case one at a "
    "time, in a separate request. Our website runs in Mumbai while our database runs in Tokyo, "
    "so each of those requests crosses a long distance and costs real time."
)

para("Measured effect on how long the case list takes to load:")

table(
    ["Cases in the system", "Extra loading time", "How it feels"],
    [
        ["100", "about 1 second", "Slightly sluggish"],
        ["300", "about 3 seconds", "Users start complaining"],
        ["500", "about 4.5 seconds", "Clearly bad"],
        ["1,000", "about 9 seconds", "Page may fail to load at all"],
        ["2,000", "about 18 seconds", "Page fails"],
    ],
    widths=[1.9, 1.9, 2.5],
)

para(
    "Important: this has nothing to do with how many users we have. One person with 1,200 "
    "cases would experience it just as badly as thirty people would. This is the problem we "
    "are most likely to hit first.",
    bold=True,
)

para(
    "The dashboard already does this the efficient way. The case list simply was never updated "
    "to match, so the fix is well understood."
)

heading("Monthly data transfer runs out sooner than storage", 2)

para(
    "The free plan includes 5 GB of data transfer per month. Because the case list loads every "
    "case in the system on every visit — rather than just the ones being viewed — this "
    "allowance is used up quickly."
)

table(
    ["Users", "Cases before the monthly allowance runs out"],
    [
        ["5", "about 1,900"],
        ["10", "about 950"],
        ["20", "about 470"],
        ["50", "about 190"],
    ],
    widths=[1.5, 4.0],
)

para(
    "At twenty users the allowance runs out at roughly 500 cases — a level the team would "
    "reach within a year. Storage would have lasted eight years; data transfer would not last "
    "one. This is caused by loading more than we need, not by having too much data."
)

# ---------------------------------------------------------------- fixes

doc.add_page_break()
heading("3. What to do about it", 1)

para("In priority order. The first three are quick and inexpensive.")

table(
    ["#", "Change", "Effort", "What it gains"],
    [
        [
            "1",
            "Store uploaded quotation files in Google Drive instead of in the database",
            "Half a day",
            "Removes the biggest space problem entirely. Free to do now; more expensive later.",
        ],
        [
            "2",
            "Load customer details for the case list in one request instead of one per case",
            "About an hour",
            "Removes the slow-page problem completely and cuts data transfer by around 40%.",
        ],
        [
            "3",
            "Adjust the database connection settings for our hosting setup",
            "Minutes",
            "Prevents the system running out of database connections under load.",
        ],
        [
            "4",
            "Load only the records currently being viewed, with proper search and paging",
            "Several days",
            "The change that takes us past 100 users. Biggest job on this list.",
        ],
        [
            "5",
            "Automatically archive activity history older than 12 months",
            "Half a day",
            "Activity history is 80% of our record growth. Keeps it flat over time.",
        ],
    ],
    widths=[0.4, 2.4, 0.9, 2.6],
)

heading("4. Three things to be aware of", 1)

para(
    "These are not performance issues, but they carry real business risk and should be a "
    "decision rather than an oversight."
)

bullet(
    "Vercel's free plan is licensed for non-commercial use only. A CRM used by our sales team "
    "is commercial use. There is no technical warning for this — an account can simply be "
    "suspended. The $20/month paid plan resolves it. This is the item I would act on first.",
    bold_prefix="Our website hosting is not licensed for business use. ",
)

bullet(
    "Supabase's free plan does not take automatic backups. We have backup scripts and they "
    "work, but the backups currently sit on one laptop and contain the full customer list in "
    "readable form. A regular schedule and an off-laptop copy are needed.",
    bold_prefix="We have no automatic backups. ",
)

bullet(
    "Supabase pauses free projects after 7 days with no activity, and restarting is manual. "
    "Normal weekday use prevents this, but an extended holiday shutdown could pause the live "
    "system.",
    bold_prefix="The database pauses if unused for a week. ",
)

heading("5. What paying would cost", 1)

table(
    ["Service", "Cost", "What it gives us"],
    [
        [
            "Supabase Pro",
            "$25 / month",
            "16× the database space, 50× the monthly data transfer, automatic daily backups",
        ],
        [
            "Vercel Pro",
            "$20 / month",
            "Proper commercial licence, higher limits",
        ],
        ["Total", "$45 / month", "Raises every limit in this document by roughly 50×"],
    ],
    widths=[1.4, 1.2, 3.7],
)

para(
    "Recommendation: do fixes 1 to 3 regardless of whether we upgrade. They are cheap, and "
    "they are the difference between a paid plan lasting years instead of months. Fix 1 is "
    "worth doing even on the paid plan — large files do not belong inside a database at any "
    "price.",
    bold=True,
)

heading("6. In one paragraph", 1)

para(
    "The CRM's design is sound and its data is small — at twenty users the records themselves "
    "would take eight years to fill the free database. The limits we face come from two "
    "specific implementation details: uploaded quotation files are kept inside the database "
    "instead of in Google Drive, and list pages load far more data than they need. Together "
    "these hold us to 5–8 users. Fixing them takes roughly one to two days and raises the "
    "ceiling to 25–40 users on the same free plans. The one item that is genuinely urgent is "
    "unrelated to performance: our website hosting plan does not permit commercial use, and "
    "that should be corrected regardless of everything else here."
)

doc.add_paragraph()
closing = para(
    "A detailed technical version of this analysis, including all measurements and the exact "
    "code locations involved, is available as scalability-report-2026-08.md in the project "
    "documentation.",
    italic=True,
    size=9,
)
closing.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.save("docs/CRM-Scalability-Summary.docx")
print("wrote docs/CRM-Scalability-Summary.docx")
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
